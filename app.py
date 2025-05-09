import streamlit as st
from docx import Document
import os
import tempfile

st.title("📁 Upload & Combine DOCX Files")

# Initialize selected files list
if 'uploaded_files' not in st.session_state:
    st.session_state.uploaded_files = []

uploaded_file = st.file_uploader("Upload a .docx file", type="docx")

if uploaded_file is not None:
    # Save to a temp file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())

    # Store file info
    st.session_state.uploaded_files.append({
        "filename": uploaded_file.name,
        "filepath": file_path
    })
    st.success(f"✅ Uploaded: {uploaded_file.name}")

# Display selected files
st.subheader("✅ Selected Files")
for i, file in enumerate(st.session_state.uploaded_files):
    col1, col2 = st.columns([4, 1])
    col1.write(file["filename"])
    if col2.button("❌ Remove", key=f"remove_{i}"):
        st.session_state.uploaded_files.pop(i)
        st.experimental_rerun()

# Generate combined docx
if st.button("Generate") and st.session_state.uploaded_files:
    output_doc = Document()
    for file in st.session_state.uploaded_files:
        try:
            sub_doc = Document(file["filepath"])
            output_doc.add_paragraph(f"--- {file['filename']} ---")
            for para in sub_doc.paragraphs:
                output_doc.add_paragraph(para.text)
        except Exception as e:
            st.error(f"Error reading {file['filename']}: {e}")

    output_path = os.path.join(tempfile.gettempdir(), "combined_output.docx")
    output_doc.save(output_path)

    st.success("🎉 Combined document generated!")
    with open(output_path, "rb") as f:
        st.download_button("📥 Download Combined DOCX", f, file_name="combined_output.docx")
